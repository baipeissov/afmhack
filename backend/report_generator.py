"""
report_generator.py — генерация официального рапорта АФМ РК в формате DOCX.

Класс AFMReportGenerator собирает по делу (case) официальный оперативный рапорт
для Агентства Республики Казахстан по финансовому мониторингу.

Запуск API:
    pip install -r requirements.txt
    uvicorn backend.report_generator:app --reload --port 8001
"""

from __future__ import annotations

import io
import os
import tempfile
from dataclasses import asdict, is_dataclass
from datetime import datetime
from typing import Any, Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# ───────────────────────────── справочники ─────────────────────────────

VIOLATION_RU = {
    "casino_betting": "Реклама нелицензированного онлайн-казино/ставок",
    "pyramid_investment": "Признаки финансовой пирамиды",
    "referral_network": "Реферальная/сетевая мошенническая схема",
    "urgency_pressure": "Психологическое давление, манипуляция срочностью",
    "hidden_engagement": "Скрытые методы вовлечения аудитории",
}

EVIDENCE_TYPE_RU = {
    "audio": "Аудио",
    "ocr": "Распознанный текст (OCR)",
    "visual": "Визуальный анализ",
    "metadata": "Метаданные",
    "probe": "Зондирование",
}

LINK_TYPE_RU = {
    "shared_telegram": "общий Telegram-канал",
    "shared_referral_link": "общую реферальную ссылку",
    "shared_hashtag": "общий хэштег",
    "shared_phone": "общий телефонный номер",
}

# уровни риска: (порог, подпись, цвет заливки HEX, цвет текста)
RISK_BANDS = [
    (0.8, "КРИТИЧЕСКИЙ РИСК", "A32D2D", "FFFFFF"),
    (0.6, "ВЫСОКИЙ РИСК", "E24B4A", "FFFFFF"),
    (0.3, "СРЕДНИЙ РИСК", "EF9F27", "1C1D1F"),
    (0.0, "НИЗКИЙ РИСК", "1D9E75", "FFFFFF"),
]


def risk_band(score: float) -> tuple[str, str, str]:
    for threshold, label, fill, text in RISK_BANDS:
        if score >= threshold:
            return label, fill, text
    return "НИЗКИЙ РИСК", "1D9E75", "FFFFFF"


def recommendation_for(score: float) -> tuple[str, str]:
    """Возвращает (рекомендация, обоснование)."""
    if score >= 0.8:
        return ("Блокировка", "Критический уровень риска — требуется немедленная блокировка ресурса и аккаунта.")
    if score >= 0.6:
        return ("Расследование", "Высокий уровень риска — передать материалы для процессуального расследования.")
    return ("Мониторинг", "Уровень риска ниже порога эскалации — поставить аккаунт на наблюдение.")


def _as_dict(obj: Any) -> Optional[dict]:
    """Нормализует ProbeResult (dataclass) или dict в обычный dict."""
    if obj is None:
        return None
    if isinstance(obj, dict):
        return obj
    if hasattr(obj, "to_dict") and callable(obj.to_dict):
        return obj.to_dict()
    if is_dataclass(obj):
        return asdict(obj)
    return dict(getattr(obj, "__dict__", {}))


# ───────────────────────────── docx-хелперы ─────────────────────────────

def _shade(cell, hex_fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_base_font(doc: Document, name: str = "Times New Roman", size: int = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)
    # гарантируем шрифт и для кириллицы (w:cs / w:eastAsia)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


# ───────────────────────────── генератор ─────────────────────────────

class AFMReportGenerator:
    def generate(self, case: dict) -> bytes:
        doc = Document()
        _set_base_font(doc)

        self._header(doc, case)
        self._section_subject(doc, case)
        self._section_risk(doc, case)
        self._section_evidence(doc, case)
        self._section_probe(doc, case)
        self._section_network(doc, case)
        self._section_recommendation(doc, case)
        self._signature(doc, case)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ── 1. Шапка ───────────────────────────────────────────────────
    def _header(self, doc: Document, case: dict) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("АГЕНТСТВО РЕСПУБЛИКИ КАЗАХСТАН\nПО ФИНАНСОВОМУ МОНИТОРИНГУ")
        r.bold = True
        r.font.size = Pt(14)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"ОПЕРАТИВНЫЙ РАПОРТ № {case.get('case_id', '—')}")
        r.bold = True
        r.font.size = Pt(13)

        created = case.get("created_at")
        created_str = self._fmt_dt(created)
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"Дата составления: {created_str}    |    Аналитик: {case.get('analyst_name', '—')}")

        self._divider(doc)

    # ── 2. Объект проверки ─────────────────────────────────────────
    def _section_subject(self, doc: Document, case: dict) -> None:
        self._heading(doc, "Раздел 1. Объект проверки")
        acc = case.get("account", {}) or {}
        rows = [
            ("Платформа", acc.get("platform", "—")),
            ("Аккаунт (handle)", f"@{acc.get('handle', '—')}"),
            ("URL", acc.get("url", "—")),
            ("Описание профиля (bio)", acc.get("bio", "—")),
        ]
        if acc.get("created_at"):
            rows.insert(2, ("Дата регистрации", self._fmt_dt(acc.get("created_at"))))
        self._kv_table(doc, rows)

    # ── 3. Итоговая оценка риска ────────────────────────────────────
    def _section_risk(self, doc: Document, case: dict) -> None:
        self._heading(doc, "Раздел 2. Итоговая оценка риска")
        score = float(case.get("risk_score", 0.0))
        label, fill, text_color = risk_band(score)

        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        _shade(cell, fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"RISK SCORE: {score:.2f}  —  {label}")
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor.from_string(text_color)
        # немного «воздуха» внутри блока
        cell.paragraphs[0].paragraph_format.space_before = Pt(10)
        cell.paragraphs[0].paragraph_format.space_after = Pt(10)

        vc = case.get("violation_class", "")
        vc_ru = VIOLATION_RU.get(vc, vc or "—")
        p = doc.add_paragraph()
        p.add_run("Класс нарушения: ").bold = True
        p.add_run(vc_ru)

    # ── 4. Доказательная база ──────────────────────────────────────
    def _section_evidence(self, doc: Document, case: dict) -> None:
        self._heading(doc, "Раздел 3. Доказательная база")
        evidence = case.get("evidence") or []
        if not evidence:
            doc.add_paragraph("Доказательные элементы не зафиксированы.")
            return

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = ["Время", "Тип сигнала", "Описание", "Уверенность"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            _shade(cell, "F1F2F4")
            run = cell.paragraphs[0].add_run(h)
            run.bold = True

        for ev in evidence:
            cells = table.add_row().cells
            cells[0].text = str(ev.get("timestamp", "—"))
            cells[1].text = EVIDENCE_TYPE_RU.get(ev.get("type", ""), ev.get("type", "—"))
            cells[2].text = str(ev.get("description", "—"))
            conf = ev.get("confidence")
            cells[3].text = f"{float(conf) * 100:.0f}%" if conf is not None else "—"

    # ── 5. Результаты зондирования ─────────────────────────────────
    def _section_probe(self, doc: Document, case: dict) -> None:
        probe = _as_dict(case.get("probe_result"))
        if not probe:
            return
        self._heading(doc, "Раздел 4. Результаты зондирования")

        url = probe.get("bio_url") or probe.get("url") or "—"
        self._kv_table(
            doc,
            [
                ("Проверенный URL", url),
                ("Заголовок лендинга", probe.get("landing_page_title") or "—"),
            ],
        )

        self._bullet_block(doc, "Обнаруженные обещания доходности:",
                           probe.get("extracted_promises") or [], empty="не обнаружены")
        self._bullet_block(doc, "Найденные контакты:",
                           probe.get("extracted_contacts") or [], empty="не обнаружены")
        self._bullet_block(doc, "Telegram-каналы:",
                           probe.get("telegram_links") or [], empty="не обнаружены")

        sha = probe.get("screenshot_sha256")
        if sha:
            p = doc.add_paragraph()
            p.add_run("Хэш скриншота-доказательства (SHA256): ").bold = True
            r = p.add_run(sha)
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            doc.add_paragraph(
                "Целостность изображения подтверждается хэшем: любое изменение файла "
                "приведёт к изменению хэша."
            ).italic = True

    # ── 6. Сетевые связи ───────────────────────────────────────────
    def _section_network(self, doc: Document, case: dict) -> None:
        conns = case.get("network_connections")
        if not conns:
            return
        self._heading(doc, "Раздел 5. Сетевые связи")

        # сгруппировать по типу связи для сводки
        by_type: dict[str, int] = {}
        for c in conns:
            lt = c.get("link_type", "—")
            by_type[lt] = by_type.get(lt, 0) + 1
        summary = ", ".join(
            f"{n} — через {LINK_TYPE_RU.get(lt, lt)}" for lt, n in by_type.items()
        )
        doc.add_paragraph(
            f"Аккаунт связан с {len(conns)} другими аккаунтами ({summary})."
        )

        for c in conns:
            handle = c.get("handle") or c.get("target") or "—"
            lt = LINK_TYPE_RU.get(c.get("link_type", ""), c.get("link_type", "связь"))
            strength = c.get("strength")
            extra = f" (сила связи {float(strength):.2f})" if strength is not None else ""
            doc.add_paragraph(f"@{handle} — через {lt}{extra}", style="List Bullet")

    # ── 7. Рекомендация ────────────────────────────────────────────
    def _section_recommendation(self, doc: Document, case: dict) -> None:
        self._heading(doc, "Раздел 6. Рекомендация")
        score = float(case.get("risk_score", 0.0))
        action, rationale = recommendation_for(score)
        p = doc.add_paragraph()
        r = p.add_run(f"Рекомендуется: {action}")
        r.bold = True
        r.font.size = Pt(13)
        _, fill, _ = risk_band(score)
        r.font.color.rgb = RGBColor.from_string(fill)
        doc.add_paragraph(rationale)

    # ── 8. Подпись ─────────────────────────────────────────────────
    def _signature(self, doc: Document, case: dict) -> None:
        self._divider(doc)
        doc.add_paragraph()
        analyst = case.get("analyst_name", "—")
        created_str = self._fmt_dt(case.get("created_at"))

        p = doc.add_paragraph()
        p.add_run("Аналитик: ").bold = True
        p.add_run(f"______________________ / {analyst} /")

        doc.add_paragraph(f"Дата: {created_str}")
        doc.add_paragraph()
        doc.add_paragraph("М.П. (место для печати)")
        note = doc.add_paragraph(
            "Документ сформирован автоматизированной системой финансового мониторинга. "
            "Все доказательства получены с публично доступных ресурсов."
        )
        note.runs[0].italic = True
        note.runs[0].font.size = Pt(8)

    # ── общие хелперы оформления ───────────────────────────────────
    def _heading(self, doc: Document, text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor.from_string("1B57CF")

    def _divider(self, doc: Document) -> None:
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "888888")
        pbdr.append(bottom)
        pPr.append(pbdr)

    def _kv_table(self, doc: Document, rows: list[tuple[str, str]]) -> None:
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for key, val in rows:
            cells = table.add_row().cells
            run = cells[0].paragraphs[0].add_run(key)
            run.bold = True
            _shade(cells[0], "F7F8F9")
            cells[1].text = str(val)
        # узкая колонка под ключи
        for row in table.rows:
            row.cells[0].width = Pt(150)

    def _bullet_block(self, doc: Document, title: str, items: list, empty: str) -> None:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        if not items:
            doc.add_paragraph(empty, style="List Bullet")
            return
        for it in items:
            doc.add_paragraph(str(it), style="List Bullet")

    @staticmethod
    def _fmt_dt(value: Any) -> str:
        if value is None:
            return "—"
        if isinstance(value, datetime):
            return value.strftime("%d.%m.%Y %H:%M")
        return str(value)


# ───────────────────────────── FastAPI ─────────────────────────────

from fastapi import FastAPI, HTTPException  # noqa: E402
from fastapi.responses import FileResponse  # noqa: E402
from starlette.background import BackgroundTask  # noqa: E402
from pydantic import BaseModel  # noqa: E402

app = FastAPI(title="АФМ · Report Generator", version="1.0.0")
_generator = AFMReportGenerator()

# Демо-хранилище дел. В проде — выборка из БД по case_id.
CASES: dict[str, dict] = {
    "CASE-2041": {
        "case_id": "CASE-2041",
        "created_at": datetime(2026, 6, 24, 10, 30),
        "analyst_name": "Айгерим Нурланова",
        "account": {
            "handle": "easy_earn_kz",
            "platform": "TikTok",
            "url": "https://tiktok.com/@easy_earn_kz",
            "bio": "Пассивный доход 30% в месяц 💸 Гарантия вывода. t.me/quick_profit_kz",
        },
        "risk_score": 0.92,
        "violation_class": "pyramid_investment",
        "evidence": [
            {"timestamp": "00:00:12", "type": "audio", "description": "Обещание гарантированного дохода 30%/мес", "confidence": 0.94, "screenshot_path": None},
            {"timestamp": "00:00:31", "type": "ocr", "description": "Текст на экране: «Гарантия вывода средств»", "confidence": 0.88, "screenshot_path": None},
            {"timestamp": "—", "type": "probe", "description": "Лендинг содержит реквизиты и Telegram-канал", "confidence": 0.9, "screenshot_path": "evidence/easy_earn_kz_x.png"},
        ],
        "probe_result": {
            "bio_url": "https://example.com/easy-earn",
            "landing_page_title": "Инвестируй и удвой капитал",
            "extracted_promises": ["30 %", "удвоим ваш депозит за 14 дней"],
            "extracted_contacts": ["+7 (701) 123-45-67", "support@easy-earn.example"],
            "telegram_links": ["https://t.me/quick_profit_kz"],
            "screenshot_sha256": "3a7bd3e2360a3d29eea436fcfb7e44c735d117c42d1c1835420b6b9942dd4f1b",
        },
        "network_connections": [
            {"handle": "casino_win_astana", "link_type": "shared_telegram", "strength": 0.9},
            {"handle": "bet_master_kz", "link_type": "shared_telegram", "strength": 0.85},
            {"handle": "referral_king_kz", "link_type": "shared_referral_link", "strength": 0.6},
        ],
    }
}


class ReportRequest(BaseModel):
    case_id: str


@app.post("/generate-report")
async def generate_report(body: ReportRequest):
    case = CASES.get(body.case_id)
    if case is None:
        raise HTTPException(status_code=404, detail=f"Дело не найдено: {body.case_id}")

    data = _generator.generate(case)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.write(data)
    tmp.close()

    filename = f"AFM_report_{body.case_id}.docx"
    return FileResponse(
        tmp.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
        background=BackgroundTask(lambda: os.unlink(tmp.name)),
    )


@app.get("/health")
async def health():
    return {"status": "ok", "cases": list(CASES.keys())}
